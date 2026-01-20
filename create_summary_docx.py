#!/usr/bin/env python3
"""
Create comprehensive CHEM 361 summary document combining:
- Data-driven methodology
- Hub-based curriculum structure
- Alignment with existing syllabus
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading(doc, text, level=1):
    """add heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows):
    """add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = str(cell_text)

    return table

def create_summary_document():
    doc = Document()

    # title
    title = doc.add_heading('CHEM 361: Data-Driven Curriculum Design', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Knowledge Graph Analysis of Inorganic Chemistry Textbooks')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph(f'McNeese State University | January 2026')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # executive summary
    add_heading(doc, 'Executive Summary', 1)

    doc.add_paragraph(
        'This document presents a data-driven approach to curriculum design for CHEM 361 '
        'Inorganic Chemistry. By analyzing prerequisite relationships extracted from 7 '
        'inorganic chemistry textbooks (8,756 text chunks), we identified the natural '
        'structure of inorganic chemistry knowledge:'
    )

    bullets = [
        '32 actionable foundations (concepts with no IC prerequisites)',
        '13 hub concepts (bottlenecks where knowledge flows through)',
        '5 major capstones (integration endpoints)',
    ]
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')

    doc.add_paragraph(
        'The resulting curriculum places foundations first, routes through hubs in '
        'prerequisite order, and concludes with capstones for integration.'
    )

    # part 1: methodology
    add_heading(doc, 'Part 1: The Data-Driven Methodology', 1)

    add_heading(doc, 'Why Data-Driven?', 2)
    doc.add_paragraph(
        'Traditional curricula are based on textbook chapter order (author preference), '
        'historical tradition, and instructor intuition. This leads to hidden connections, '
        'implicit prerequisites, and misidentified "foundations" that are actually capstones.'
    )
    doc.add_paragraph('Our approach: Let the data reveal the structure.')

    add_heading(doc, 'Knowledge Graph Construction', 2)
    doc.add_paragraph('Input: 7 IC textbooks → 8,756 text chunks')
    doc.add_paragraph('Extraction: For each chunk, LLM extracted:')
    bullets = [
        'Main topic',
        'Prerequisites ("what must student know first?")',
        'Leads-to ("what does this enable?")',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph('Output: Directed graph with 5,374 nodes and 2,885 edges')

    add_heading(doc, 'Graph Analysis', 2)
    headers = ['Metric', 'Value', 'Implication']
    rows = [
        ['Nodes', '5,374', 'Rich concept vocabulary'],
        ['Edges', '2,885', 'Prerequisite relationships'],
        ['Mean degree', '1.05', 'Sparse (tree-like) structure'],
        ['In-degree = 0', '86.9%', 'Most concepts are entry points'],
        ['Out-degree = 0', '82.4%', 'Most concepts are endpoints'],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    doc.add_paragraph(
        'Key insight: The graph is sparse. PageRank is inappropriate for sparse graphs. '
        'Degree-based analysis is the correct approach.'
    )

    add_heading(doc, 'Classification Method', 2)
    doc.add_paragraph('FOUNDATION: in_degree = 0, out_degree ≥ 5')
    doc.add_paragraph('    → No IC prerequisites, enables many topics → TEACH FIRST', style='List Bullet')
    doc.add_paragraph('HUB: in_degree > 2 AND out_degree > 2')
    doc.add_paragraph('    → Knowledge flows THROUGH these → TEACH IN MIDDLE', style='List Bullet')
    doc.add_paragraph('CAPSTONE: out_degree = 0, in_degree ≥ 50')
    doc.add_paragraph('    → Many prerequisites converge here → TEACH LAST', style='List Bullet')

    # part 2: the 13 hubs
    add_heading(doc, 'Part 2: The 13 Hubs (Knowledge Bottlenecks)', 1)

    doc.add_paragraph(
        'Every prerequisite chain passes through at least one of these 13 hubs. '
        'They are the critical path of the curriculum.'
    )

    headers = ['Rank', 'Hub', 'In', 'Out', 'Role', 'Teach When']
    rows = [
        ['1', 'Acid-Base Chemistry', '55', '7', 'Connector to applications', 'Week 6-7'],
        ['2', 'Crystal Field Theory', '46', '12', 'Electronic structure gateway', 'Week 10-11'],
        ['3', 'Molecular Orbital Theory', '43', '13', 'Bonding theory gateway', 'Week 8-9'],
        ['4', 'Redox Chemistry', '7', '34', 'Reaction chemistry gateway', 'Week 5'],
        ['5', 'Periodic Trends', '17', '19', 'Descriptive chemistry gateway', 'Week 2-3'],
        ['6', 'Chemical Bonding', '32', '3', 'Foundation receiver', 'Week 4'],
        ['7', 'Organometallic Chemistry', '24', '5', 'Applications gateway', 'Week 12'],
        ['8', 'Atomic Structure', '13', '12', 'Quantum-electronic bridge', 'Week 1-2'],
        ['9', 'Transition Metal Chemistry', '11', '7', 'd-block integration', 'Week 13-14'],
        ['10', 'Crystal Structures', '5', '7', 'Solid state gateway', 'Week 9'],
        ['11', 'Thermochemistry', '7', '3', 'Energetics gateway', 'Week 4-5'],
        ['12', 'Electronegativity', '3', '4', 'Polarity concepts', 'Week 3'],
        ['13', 'Polymer Chemistry', '4', '3', 'Materials gateway', 'Week 15'],
    ]
    add_table(doc, headers, rows)

    # part 3: curriculum structure
    add_heading(doc, 'Part 3: The Curriculum Structure', 1)

    add_heading(doc, 'Phase Overview', 2)
    doc.add_paragraph('PHASE 1: FOUNDATIONS (Weeks 1-4)')
    doc.add_paragraph('    Build the base. No IC prerequisites needed.', style='List Bullet')
    doc.add_paragraph('PHASE 2: HUB TRAVERSAL (Weeks 5-12)')
    doc.add_paragraph('    Navigate through the 13 hubs. Each hub unlocks the next.', style='List Bullet')
    doc.add_paragraph('PHASE 3: CAPSTONES (Weeks 13-16)')
    doc.add_paragraph('    Integration and application. All hubs converge here.', style='List Bullet')

    add_heading(doc, 'Phase 1: Foundations (Weeks 1-4)', 2)

    doc.add_paragraph('Week 1: Atomic Foundations')
    headers = ['Topic', 'Out-degree', 'Why First']
    rows = [
        ['Electron Configuration', '25', 'Enables CFT, MO theory, periodic trends'],
        ['Quantum Numbers', '10', 'Enables orbital understanding'],
        ['Atomic Orbitals', '7', 'Enables bonding theories'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    doc.add_paragraph('Week 2: Periodic Foundations')
    headers = ['Topic', 'Out-degree', 'Why Now']
    rows = [
        ['Periodic Table Trends', '7', 'Enables element chemistry'],
        ['Electronegativity basics', '4', 'Enables bonding predictions'],
        ['Ionization Energy', '4', 'Enables reactivity predictions'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    doc.add_paragraph('Week 3: Bonding Foundations')
    headers = ['Topic', 'Out-degree', 'Why Now']
    rows = [
        ['Ionic Bonding', '16', 'Enables crystal structures'],
        ['Coordination Fundamentals', '40', 'Enables all coordination topics'],
        ['Lewis Structures', '6', 'Enables acid-base theory'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    doc.add_paragraph('Week 4: Energetics Foundations')
    headers = ['Topic', 'Out-degree', 'Why Now']
    rows = [
        ['Oxidation States', '17', 'Enables redox chemistry'],
        ['Thermodynamics basics', '9', 'Enables stability predictions'],
        ['Lattice Energy concepts', '5', 'Enables ionic compound chemistry'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    add_heading(doc, 'Phase 2: Hub Traversal (Weeks 5-12)', 2)

    doc.add_paragraph('Week 5: Redox Hub')
    doc.add_paragraph('Prerequisites: Oxidation States (Week 4), Electron Configuration (Week 1)')
    doc.add_paragraph('Topics: Electrochemical Series, Standard Potentials, Nernst Equation')
    doc.add_paragraph('Unlocks: Electrochemistry applications, Main Group reactions, TM chemistry')
    doc.add_paragraph()

    doc.add_paragraph('Weeks 6-7: Acid-Base Hub')
    doc.add_paragraph('Prerequisites: Periodic Trends (Week 2), Lewis Structures (Week 3)')
    doc.add_paragraph('Topics: Brønsted-Lowry, Lewis Acid-Base, HSAB Principle, pKa Trends')
    doc.add_paragraph('Unlocks: Coordination chemistry, Main Group reactions, Catalysis')
    doc.add_paragraph()

    doc.add_paragraph('Weeks 8-9: Molecular Orbital Hub')
    doc.add_paragraph('Prerequisites: Atomic Orbitals (Week 1), Chemical Bonding (Week 3)')
    doc.add_paragraph('Topics: LCAO Approach, Bonding/Antibonding, MO Diagrams, Bond Order')
    doc.add_paragraph('Unlocks: Crystal Field Theory, Main Group bonding, Solid State band theory')
    doc.add_paragraph()

    doc.add_paragraph('Weeks 10-11: Crystal Field Theory Hub')
    doc.add_paragraph('Prerequisites: Electron Config (Week 1), Coord. Fundamentals (Week 3), MO Theory (Week 8-9)')
    doc.add_paragraph('Topics: d-Orbital Splitting, Spectrochemical Series, CFSE, High/Low Spin, Jahn-Teller')
    doc.add_paragraph('This is where "Why is Cu²⁺ blue?" gets answered.')
    doc.add_paragraph()

    doc.add_paragraph('Week 12: Organometallic Hub')
    doc.add_paragraph('Prerequisites: MO Theory (Week 8-9), CFT (Week 10-11)')
    doc.add_paragraph('Topics: Metal-Carbon Bonds, 18-Electron Rule, Oxidative Addition, Carbonyls')
    doc.add_paragraph('Unlocks: Catalysis applications, Industrial chemistry')
    doc.add_paragraph()

    add_heading(doc, 'Phase 3: Capstones (Weeks 13-16)', 2)

    doc.add_paragraph('Weeks 13-14: Transition Metal Chemistry Capstone')
    doc.add_paragraph('Prerequisites converging: CFT, Redox, Coordination Fundamentals, Periodic Trends')
    doc.add_paragraph('Topics: d-Block Survey, Oxidation State Patterns, Biological Relevance')
    doc.add_paragraph()

    doc.add_paragraph('Week 15: Coordination Chemistry Capstone')
    doc.add_paragraph('Prerequisites converging: All bonding theories, CFT, Acid-Base, Redox')
    doc.add_paragraph('Topics: Complex Synthesis, Isomerism, Reaction Mechanisms, Applications')
    doc.add_paragraph()

    doc.add_paragraph('Week 16: Main Group Chemistry Capstone')
    doc.add_paragraph('Prerequisites converging: Periodic Trends, All bonding theories, Acid-Base, Redox')
    doc.add_paragraph('Why Main Group is LAST: in_degree=368 (most prerequisites), out_degree=0 (nothing depends on it)')
    doc.add_paragraph()

    # part 4: question mapping
    add_heading(doc, 'Part 4: Question-to-Curriculum Mapping', 1)

    doc.add_paragraph('Example: "Why is copper sulfate blue?"')
    doc.add_paragraph()
    doc.add_paragraph('Trace through curriculum:')
    doc.add_paragraph('Question: Why is Cu²⁺ blue?', style='List Bullet')
    doc.add_paragraph('→ CAPSTONE: Coordination Chemistry (Week 15)', style='List Bullet')
    doc.add_paragraph('→ HUB: Crystal Field Theory (Week 10-11)', style='List Bullet')
    doc.add_paragraph('→ HUB: MO Theory (Week 8-9)', style='List Bullet')
    doc.add_paragraph('→ FOUNDATION: Electron Configuration (Week 1)', style='List Bullet')
    doc.add_paragraph()

    doc.add_paragraph('Answer construction:')
    doc.add_paragraph('1. Cu²⁺ has electron configuration [Ar] 3d⁹ (Week 1)')
    doc.add_paragraph('2. In octahedral field, d-orbitals split (Week 10)')
    doc.add_paragraph('3. Energy gap Δ corresponds to orange light (~600 nm) (Week 10)')
    doc.add_paragraph('4. Absorbs orange → transmits blue (Week 10-11)')

    add_heading(doc, 'More Question Examples', 2)
    headers = ['Question', 'Foundation', 'Hub(s)', 'Capstone']
    rows = [
        ['Why is Au golden?', 'Electron Config', 'MO Theory', 'Transition Metals'],
        ['How do batteries work?', 'Oxidation States', 'Redox Chemistry', 'Electrochemistry'],
        ['Why is O₂ paramagnetic?', 'Atomic Orbitals', 'MO Theory', 'Main Group'],
        ['How do enzymes use metals?', 'Coord. Fundamentals', 'CFT, Redox', 'Bioinorganic'],
    ]
    add_table(doc, headers, rows)

    # part 5: alignment with syllabus
    add_heading(doc, 'Part 5: Alignment with CHEM 361 Syllabus', 1)

    add_heading(doc, 'SLO Mapping', 2)
    headers = ['SLO', 'Description', 'Hub(s) Supporting']
    rows = [
        ['1', 'Explain periodic trends', 'Periodic Trends (Hub 5), Electronegativity (Hub 12)'],
        ['2', 'Predict geometry, magnetism, color', 'CFT (Hub 2), MO Theory (Hub 3)'],
        ['3', 'Apply bonding theories', 'Chemical Bonding (Hub 6), MO Theory (Hub 3)'],
        ['4', 'Analyze main group reactivity', 'Acid-Base (Hub 1), Redox (Hub 4)'],
        ['5', 'Interpret solid state properties', 'Crystal Structures (Hub 10)'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    add_heading(doc, 'Assessment Alignment', 2)
    doc.add_paragraph('Midterm 1 (Coordination): Tests Hubs 2, 3, 6 - CFT, MO Theory, Bonding')
    doc.add_paragraph('Midterm 2 (Main Group & Solid): Tests Hubs 1, 4, 5, 10 - Acid-Base, Redox, Periodic, Crystal')
    doc.add_paragraph('Final Exam: Integration across all 13 hubs and 5 capstones')
    doc.add_paragraph()

    add_heading(doc, 'Hub Checkpoints', 2)
    headers = ['Hub', 'Checkpoint Question', 'Must Master Before']
    rows = [
        ['Redox', 'Balance redox in acidic solution', 'Acid-Base Hub'],
        ['Acid-Base', 'Predict acidity using HSAB', 'MO Theory Hub'],
        ['MO Theory', 'Draw MO diagram for CO', 'CFT Hub'],
        ['CFT', 'Calculate CFSE high-spin vs low-spin', 'Organometallic Hub'],
    ]
    add_table(doc, headers, rows)

    # part 6: data summary
    add_heading(doc, 'Part 6: Data Summary', 1)

    add_heading(doc, 'Graph Statistics', 2)
    headers = ['Metric', 'Value']
    rows = [
        ['Textbooks analyzed', '7'],
        ['Text chunks', '8,756'],
        ['Total nodes', '5,374'],
        ['Prerequisite edges', '2,885'],
        ['Mean degree', '1.05'],
        ['Actionable foundations', '32'],
        ['Hubs identified', '13'],
        ['Major capstones', '5'],
    ]
    add_table(doc, headers, rows)
    doc.add_paragraph()

    add_heading(doc, 'Capstone In-Degrees', 2)
    headers = ['Capstone', 'In-degree', 'Coverage']
    rows = [
        ['Main Group Chemistry', '368', '40% of chunks'],
        ['Coordination Chemistry', '157', '43% of chunks'],
        ['Solid State Chemistry', '107', '17% of chunks'],
        ['Bioinorganic Chemistry', '84', '8% of chunks'],
        ['Transition Metal Chemistry', '72', '15% of chunks'],
    ]
    add_table(doc, headers, rows)

    # conclusion
    add_heading(doc, 'Conclusion', 1)

    doc.add_paragraph(
        'This data-driven curriculum represents a fundamental shift from traditional '
        'topic ordering. By letting the prerequisite relationships in 7 textbooks reveal '
        'the natural structure of inorganic chemistry knowledge, we identified 13 hub '
        'concepts that serve as knowledge bottlenecks.'
    )
    doc.add_paragraph(
        'The key insight: Main Group Chemistry (traditionally taught early) has the '
        'highest in-degree (368 prerequisites) and zero out-degree. It is the ultimate '
        'integration endpoint, not a foundation. Teaching it last allows students to '
        'truly integrate all prior knowledge.'
    )
    doc.add_paragraph(
        'Students following this curriculum will always know "why they need to learn '
        'this" because every concept is explicitly connected to what comes before and '
        'after through the hub structure.'
    )

    # references
    add_heading(doc, 'References and Resources', 1)
    doc.add_paragraph('Knowledge graph: experiments/results/knowledge_graph.json')
    doc.add_paragraph('Hub analysis: experiments/results/hub_analysis.json')
    doc.add_paragraph('Full hub documentation: docs/HUB_CURRICULUM.md')
    doc.add_paragraph('Data-driven curriculum: docs/DATA_DRIVEN_CURRICULUM.md')

    doc.add_paragraph()
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %Y")}')
    doc.add_paragraph('Analysis: 7 textbooks, 8,756 chunks, 5,374 nodes, 2,885 edges')

    # save
    output_path = '/home/kiran/CHEM361_DataDriven_Summary.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_summary_document()
